import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document
import io

# Cấu hình đường dẫn Tesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Mở PDF
pdf_path = "file.pdf"
docx_path = "output.docx"

doc_pdf = fitz.open(pdf_path)
doc_word = Document()

# Tùy chọn: tiêu đề văn bản
doc_word.add_heading("Kết quả OCR từ PDF", level=1)

for i, page in enumerate(doc_pdf):
    # Chuyển mỗi trang thành ảnh PNG
    pix = page.get_pixmap(dpi=300)
    img_bytes = pix.tobytes("png")
    image = Image.open(io.BytesIO(img_bytes))

    # OCR từng ảnh (lang='vie+eng' nếu cần)
    text = pytesseract.image_to_string(image, lang='vie+eng')

    # Tách theo dòng để tái tạo đoạn văn
    for para in text.strip().split("\n\n"):
        if para.strip():  # loại dòng trắng
            doc_word.add_paragraph(para.strip())

    # Trang mới
    doc_word.add_page_break()

# Lưu file Word
doc_word.save(docx_path)
print(f"✅ Đã chuyển {pdf_path} thành {docx_path} với định dạng cơ bản.")
